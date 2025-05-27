
import streamlit as st
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import requests
from bs4 import BeautifulSoup
import os
from PIL import Image
from datetime import datetime, timedelta

# Setup logging
logging.basicConfig(filename="logs.txt", level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Sample project data
sample_data = {
    "title": "Youth Empowerment Through Dance",
    "overview": "This project aims to empower youth in Mitchells Plain through dance, fostering creativity, discipline, and self-expression.",
    "objectives": [
        "Provide dance training to 100 youth aged 10–18.",
        "Host two community dance showcases.",
        "Train 10 local youth leaders in arts facilitation."
    ],
    "impact": "Participants will gain confidence, leadership skills, and a positive community identity.",
    "budget_total": "120000",
    "budget_items": {
        "Instructor fees": "40000",
        "Venue rental": "20000",
        "Costumes and props": "25000",
        "Marketing": "15000",
        "Administration": "20000"
    },
    "timeline": "June to December 2025",
    "demographics": {
        "description": "Targeting disadvantaged youth in Tafelsig, primarily between ages 10–18, majority Coloured and Black communities."
    }
}

def generate_proposal(project_data, logo_path=None):
    try:
        doc = Document()
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "Ubuntu Dance and Theatre Arts\nTafelsig, Mitchells Plain, Western Cape, South Africa\nEmail: info@ubuntudance.org | Phone: +27 21 123 4567"
        header_para.style.font.size = Pt(10)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("[Ubuntu Dance and Theatre Arts Logo]", style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading(f"{project_data['title']} Grant Proposal", 0)
        doc.add_paragraph(f"Prepared for: Potential Funders\nDate: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_page_break()
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(project_data['overview'])
        doc.add_heading("Objectives", level=1)
        for obj in project_data['objectives']:
            doc.add_paragraph(f"- {obj}", style="List Bullet")
        doc.add_heading("Community Impact", level=1)
        doc.add_paragraph(project_data['impact'])
        doc.add_heading("Budget", level=1)
        doc.add_paragraph(f"Total Budget: {project_data['budget_total']} ZAR")
        for item, amount in project_data['budget_items'].items():
            doc.add_paragraph(f"{item}: {amount} ZAR")
        doc.add_heading("Timeline", level=1)
        doc.add_paragraph(project_data['timeline'])
        doc.add_heading("Appendices", level=1)
        doc.add_heading("Demographics", level=2)
        doc.add_paragraph(project_data['demographics']['description'])
        doc.add_heading("Team", level=2)
        doc.add_paragraph("Project Lead: [Your Name]\nArtistic Director: [Director Name]\nCommunity Liaison: [Liaison Name]")
        filename = os.path.join(os.getcwd(), f"{project_data['title'].replace(' ', '_')}_Proposal.docx")
        doc.save(filename)
        return filename
    except Exception as e:
        logging.exception("Proposal generation failed")
        st.error(f"Failed to generate proposal: {e}")
        return None

def find_grants(keywords, cache_file="grants_cache.json", expiration_hours=24):
    grants = []
    if os.path.exists(cache_file):
        try:
            with open(cache_file, "r", encoding="utf-8") as f:
                cached = json.load(f)
                timestamp = datetime.strptime(cached.get("last_updated", "1970-01-01 00:00:00"), "%Y-%m-%d %H:%M:%S")
                if datetime.now() - timestamp < timedelta(hours=expiration_hours):
                    st.info("Using cached grant data.")
                    return cached.get("data", [])
        except Exception as e:
            logging.warning(f"Cache load failed: {e}")
            st.error(f"Failed to load cached grants: {e}")

    urls = ["https://www.nhc.org.za/grants", "https://www.nac.org.za/funding"]
    for url in urls:
        try:
            response = requests.get(url, timeout=5)
            soup = BeautifulSoup(response.text, 'html.parser')
            for link in soup.find_all('a', href=True):
                text = link.get_text().lower()
                if any(k.lower() in text for k in keywords):
                    grants.append({
                        "title": link.get_text(),
                        "url": link['href'],
                        "source": url,
                        "summary": text[:60] + "..." if len(text) > 60 else text,
                        "last_updated": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    })
        except Exception as e:
            logging.warning(f"Scraping failed for {url}: {e}")
            st.error(f"Error fetching grants from {url}")

    try:
        with open(cache_file, "w", encoding="utf-8") as f:
            json.dump({"last_updated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "data": grants}, f, indent=2)
    except Exception as e:
        logging.error(f"Failed to cache grants: {e}")

    return grants

st.title("Ubuntu Dance and Theatre Arts – Grant Proposal App")

if st.checkbox("Use sample project data"):
    project_data = sample_data

    if st.button("Generate Proposal Document"):
        with st.spinner("Creating proposal..."):
            doc_path = generate_proposal(project_data)
            if doc_path:
                with open(doc_path, "rb") as f:
                    st.success("Proposal generated successfully!")
                    st.download_button(
                        label="📥 Download Proposal",
                        data=f,
                        file_name=os.path.basename(doc_path),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

    st.subheader("🔍 Search for Grant Opportunities")
    keywords = st.text_input("Enter keywords (comma-separated)", "youth, dance, arts")
    if st.button("Search Grants"):
        keyword_list = [k.strip() for k in keywords.split(",")]
        results = find_grants(keyword_list)
        if results:
            for grant in results:
                st.markdown(
                    f"**[{grant['title']}]({grant['url']})**  
"
                    f"Source: {grant['source']}  
"
                    f"Summary: {grant['summary']}"
                )
        else:
            st.info("No grants found.")
else:
    st.info("Tick the checkbox above to load sample data.")
